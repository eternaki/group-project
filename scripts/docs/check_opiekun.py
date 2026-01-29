#!/usr/bin/env python3
"""Sprawdz czy Opiekun projektu zostal wypelniony."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document


def main():
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "docs" / "deliverables" / "DTP_v2.docx"

    doc = Document(docx_path)

    print("=== SZUKANIE 'Opiekun projektu' ===\n")

    if doc.tables:
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                if "opiekun" in cell.text.lower():
                    print(f"Row {row_idx}, Cell {cell_idx}:")
                    print(f"  Text: {cell.text[:100]}...")
                    print()


if __name__ == "__main__":
    main()
