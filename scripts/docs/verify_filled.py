#!/usr/bin/env python3
"""Weryfikacja wypełnionego szablonu DTP."""

from pathlib import Path
from docx import Document


def verify_template(docx_path: Path):
    """Weryfikuje zawartość wypełnionego szablonu."""
    doc = Document(docx_path)

    print("=" * 60)
    print("WERYFIKACJA WYPEŁNIONEGO SZABLONU DTP")
    print("=" * 60)

    print("\n--- TABELA NAGŁÓWKOWA ---")
    if doc.tables:
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for cell in row.cells:
                text = cell.text.strip()[:40]
                if text:
                    cells_text.append(text)
            if cells_text:
                print(f"Row {row_idx}: {' | '.join(cells_text)}")

    print("\n--- SEKCJE TREŚCI (pierwsze 100 znaków) ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 5:
            style = para.style.name if para.style else "None"
            preview = text[:100].replace('\n', ' ')
            print(f"[{i:3d}] [{style:20s}] {preview}...")

    print("\n--- TABELA ZAŁĄCZNIKÓW ---")
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for row_idx, row in enumerate(table.rows):
            cells_text = [cell.text.strip()[:30] for cell in row.cells]
            print(f"Row {row_idx}: {' | '.join(cells_text)}")


def main():
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "docs" / "deliverables" / "DTP_filled.docx"

    if not docx_path.exists():
        print(f"Błąd: Nie znaleziono pliku {docx_path}")
        return

    verify_template(docx_path)


if __name__ == "__main__":
    main()
