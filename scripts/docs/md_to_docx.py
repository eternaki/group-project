#!/usr/bin/env python3
"""
Skrypt do konwersji pliku Markdown (DTP.md) do formatu Word (.docx).

Obsługuje:
- Nagłówki (H1-H4)
- Paragrafy
- Tabele
- Bloki kodu
- Listy
- Pogrubienie i kursywę
"""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str):
    """Ustawia kolor tła komórki tabeli."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_code_block(doc: Document, code: str, language: str = ""):
    """Dodaje blok kodu z szarym tłem."""
    # Tworzymy tabelę 1x1 dla bloku kodu
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Szare tło
    set_cell_shading(cell, "E8E8E8")

    # Tekst kodu
    para = cell.paragraphs[0]
    run = para.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # Dodaj pustą linię po bloku
    doc.add_paragraph()


def parse_table(lines: list[str]) -> list[list[str]]:
    """Parsuje tabelę markdown do listy list."""
    rows = []
    for line in lines:
        if line.strip().startswith('|') and '---' not in line:
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    """Dodaje tabelę do dokumentu Word."""
    if not rows:
        return

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = cell_text.strip()

            # Nagłówek tabeli - pogrubiony z szarym tłem
            if i == 0:
                set_cell_shading(cell, "D9D9D9")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    doc.add_paragraph()


def format_inline(text: str):
    """Zwraca listę fragmentów z formatowaniem (bold, italic, code)."""
    fragments = []
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'

    last_end = 0
    for match in re.finditer(pattern, text):
        # Tekst przed dopasowaniem
        if match.start() > last_end:
            fragments.append(('normal', text[last_end:match.start()]))

        matched = match.group()
        if matched.startswith('**'):
            fragments.append(('bold', matched[2:-2]))
        elif matched.startswith('*'):
            fragments.append(('italic', matched[1:-1]))
        elif matched.startswith('`'):
            fragments.append(('code', matched[1:-1]))

        last_end = match.end()

    # Tekst po ostatnim dopasowaniu
    if last_end < len(text):
        fragments.append(('normal', text[last_end:]))

    return fragments if fragments else [('normal', text)]


def add_formatted_paragraph(doc: Document, text: str, style: str = None):
    """Dodaje paragraf z formatowaniem inline."""
    para = doc.add_paragraph(style=style)

    fragments = format_inline(text)
    for fmt, content in fragments:
        run = para.add_run(content)
        if fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True
        elif fmt == 'code':
            run.font.name = "Consolas"
            run.font.size = Pt(10)


def convert_md_to_docx(md_path: Path, docx_path: Path):
    """Konwertuje plik Markdown do Word."""

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Ustawienia dokumentu
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    i = 0

    # Pomijamy pierwszą linię jeśli to instrukcja kopiowania
    if lines and lines[0].startswith('>'):
        i = 1
        while i < len(lines) and lines[i].strip() in ['', '---']:
            i += 1

    in_code_block = False
    code_lines = []
    code_lang = ""
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Blok kodu
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lang = line.strip()[3:]
                code_lines = []
            else:
                in_code_block = False
                add_code_block(doc, '\n'.join(code_lines), code_lang)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Tabela
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            in_table = False
            rows = parse_table(table_lines)
            add_table(doc, rows)
            table_lines = []

        # Nagłówki
        if line.startswith('####'):
            doc.add_heading(line[4:].strip(), level=4)
        elif line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line[1:].strip(), level=1)

        # Linia pozioma
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)

        # Lista numerowana
        elif re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            add_formatted_paragraph(doc, text, style='List Number')

        # Lista punktowana
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            add_formatted_paragraph(doc, text, style='List Bullet')

        # Pusty paragraf
        elif not line.strip():
            pass  # Pomijamy puste linie

        # Normalny tekst
        else:
            add_formatted_paragraph(doc, line)

        i += 1

    # Zapisz jeśli tabela na końcu
    if in_table and table_lines:
        rows = parse_table(table_lines)
        add_table(doc, rows)

    doc.save(docx_path)
    print(f"Zapisano: {docx_path}")


def main():
    """Główna funkcja."""
    project_root = Path(__file__).parent.parent.parent

    md_path = project_root / "docs" / "deliverables" / "DTP.md"
    docx_path = project_root / "docs" / "deliverables" / "DTP.docx"

    if not md_path.exists():
        print(f"Błąd: Nie znaleziono pliku {md_path}")
        return

    convert_md_to_docx(md_path, docx_path)
    print(f"\nKonwersja zakończona!")
    print(f"Plik źródłowy: {md_path}")
    print(f"Plik wynikowy: {docx_path}")


if __name__ == "__main__":
    main()
