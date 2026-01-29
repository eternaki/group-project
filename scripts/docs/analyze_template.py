#!/usr/bin/env python3
"""
Skrypt do analizy struktury szablonu .docx.
"""

from pathlib import Path
from docx import Document


def analyze_template(docx_path: Path):
    """Analizuje strukturę szablonu."""
    doc = Document(docx_path)

    print("=" * 60)
    print("ANALIZA SZABLONU DTP")
    print("=" * 60)

    print("\n--- STYLE DOKUMENTU ---")
    for style in doc.styles:
        if style.type == 1:  # Paragraph styles
            if "Heading" in style.name or "Title" in style.name or "Nagłówek" in style.name:
                print(f"  {style.name}")

    print("\n--- PARAGRAFY ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else "None"
        if text:
            print(f"[{i:3d}] [{style:20s}] {text[:80]}")

    print("\n--- TABELE ---")
    for i, table in enumerate(doc.tables):
        print(f"\nTabela {i + 1}: {len(table.rows)} wierszy x {len(table.columns)} kolumn")
        for row_idx, row in enumerate(table.rows):
            row_text = " | ".join(cell.text.strip()[:20] for cell in row.cells)
            print(f"  Wiersz {row_idx}: {row_text}")

    print("\n--- SEKCJE ---")
    for i, section in enumerate(doc.sections):
        print(f"Sekcja {i + 1}:")
        print(f"  Page width: {section.page_width.inches:.2f} in")
        print(f"  Page height: {section.page_height.inches:.2f} in")
        print(f"  Margins: L={section.left_margin.inches:.2f}, R={section.right_margin.inches:.2f}")


def main():
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "template" / "PG_WETI_DTP_wer. 1.00.docx"

    if not docx_path.exists():
        print(f"Błąd: Nie znaleziono pliku {docx_path}")
        return

    analyze_template(docx_path)


if __name__ == "__main__":
    main()
