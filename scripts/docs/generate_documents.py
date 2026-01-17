#!/usr/bin/env python3
"""
Generator dokumentów Word z szablonów Markdown.

Skrypt czyta pliki Markdown z docs/deliverables/ i generuje dokumenty Word
zgodne z szablonami WETI.

Użycie:
    python scripts/docs/generate_documents.py --all
    python scripts/docs/generate_documents.py --document DPP
    python scripts/docs/generate_documents.py --document Specyfikacja --output output/
    python scripts/docs/generate_documents.py --list
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print("Błąd: Wymagana biblioteka python-docx")
    print("Zainstaluj: pip install python-docx")
    sys.exit(1)


# Konfiguracja projektu - UZUPEŁNIJ DANE ZESPOŁU
PROJECT_CONFIG = {
    "temat": "Dog FACS Dataset - Pipeline do automatycznej anotacji emocji psów z wykorzystaniem AI",
    "rok_akademicki": "2025/2026",
    "semestr": "1 (zimowy)",
    "katedra": "Katedra Systemów Multimedialnych",
    "opiekun": "dr inż. [Imię Nazwisko]",
    "zespol": [
        {"imie": "Danylo Logachev", "album": "000000", "rola": "Team Lead / ML Engineer"},
        {"imie": "Anton [Nazwisko]", "album": "000000", "rola": "ML Engineer"},
        {"imie": "Danylo Zhernosek", "album": "000000", "rola": "Data Engineer"},
        {"imie": "Mariia [Nazwisko]", "album": "000000", "rola": "QA / Annotator"},
    ],
    "repo_url": "https://github.com/eternaki/group-project",
    "sprinty": [
        {"nr": 1, "nazwa": "Project Setup", "status": "Ukończony"},
        {"nr": 2, "nazwa": "Dog Detection", "status": "Ukończony"},
        {"nr": 3, "nazwa": "Breed Classification", "status": "Ukończony"},
        {"nr": 4, "nazwa": "Keypoint Detection", "status": "Ukończony"},
        {"nr": 5, "nazwa": "Emotion Classification", "status": "Ukończony"},
        {"nr": 6, "nazwa": "Inference Pipeline", "status": "Ukończony"},
        {"nr": 7, "nazwa": "Demo Application", "status": "Ukończony"},
        {"nr": 8, "nazwa": "Data Collection", "status": "Ukończony"},
        {"nr": 9, "nazwa": "Batch Annotation", "status": "Ukończony"},
        {"nr": 10, "nazwa": "Manual Verification", "status": "Ukończony"},
        {"nr": 11, "nazwa": "Dataset Finalization", "status": "Ukończony"},
        {"nr": 12, "nazwa": "Statistics & Reporting", "status": "Ukończony"},
    ],
    # Metryki - DO UZUPEŁNIENIA po rzeczywistych testach
    "metryki": {
        "detector_map": "TBD",
        "breed_top5": "TBD",
        "keypoints_pck": "TBD",
        "emotion_f1": "TBD",
    },
    "dataset": {
        "liczba_obrazow": "TBD",
        "liczba_anotacji": "TBD",
        "unikalne_rasy": "TBD",
        "keypoints_per_pies": "TBD",
    },
}


@dataclass
class DocumentSection:
    """Reprezentuje sekcję dokumentu."""

    title: str
    level: int
    content: list[str] = field(default_factory=list)
    subsections: list["DocumentSection"] = field(default_factory=list)


class MarkdownParser:
    """Parser plików Markdown."""

    def __init__(self, content: str) -> None:
        self.content = content
        self.lines = content.split("\n")

    def parse(self) -> list[DocumentSection]:
        """Parsuje Markdown i zwraca strukturę sekcji."""
        sections: list[DocumentSection] = []
        current_section: Optional[DocumentSection] = None
        current_content: list[str] = []

        for line in self.lines:
            # Pomijaj linie z instrukcjami (> Skopiuj...)
            if line.startswith(">"):
                continue

            # Wykryj nagłówki
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                # Zapisz poprzednią sekcję
                if current_section:
                    current_section.content = current_content
                    sections.append(current_section)

                level = len(header_match.group(1))
                title = header_match.group(2)
                current_section = DocumentSection(title=title, level=level)
                current_content = []
            else:
                current_content.append(line)

        # Zapisz ostatnią sekcję
        if current_section:
            current_section.content = current_content
            sections.append(current_section)

        return sections

    def extract_tables(self) -> list[list[list[str]]]:
        """Ekstraktuje tabele z Markdown."""
        tables: list[list[list[str]]] = []
        current_table: list[list[str]] = []
        in_table = False

        for line in self.lines:
            if "|" in line and not line.strip().startswith("```"):
                # Pomijaj linię separatora (|---|---|)
                if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                    continue

                cells = [cell.strip() for cell in line.split("|")[1:-1]]
                if cells:
                    current_table.append(cells)
                    in_table = True
            elif in_table:
                if current_table:
                    tables.append(current_table)
                current_table = []
                in_table = False

        if current_table:
            tables.append(current_table)

        return tables


class WordDocumentGenerator:
    """Generator dokumentów Word."""

    def __init__(self, config: dict) -> None:
        self.config = config
        self.doc: Optional[Document] = None

    def create_document(self) -> Document:
        """Tworzy nowy dokument Word."""
        self.doc = Document()
        self._setup_styles()
        return self.doc

    def _setup_styles(self) -> None:
        """Konfiguruje style dokumentu."""
        if not self.doc:
            return

        # Styl normalny
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(12)

    def add_title(self, text: str, level: int = 1) -> None:
        """Dodaje tytuł/nagłówek."""
        if not self.doc:
            return

        if level == 1:
            self.doc.add_heading(text, level=0)
        else:
            self.doc.add_heading(text, level=min(level, 9))

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> None:
        """Dodaje akapit."""
        if not self.doc:
            return

        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic

    def add_table(
        self,
        data: list[list[str]],
        header: bool = True,
        col_widths: Optional[list[float]] = None,
    ) -> None:
        """Dodaje tabelę."""
        if not self.doc or not data:
            return

        rows = len(data)
        cols = len(data[0]) if data else 0

        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_text)

                # Pogrubienie nagłówka
                if header and i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # Ustaw szerokości kolumn
        if col_widths:
            for i, width in enumerate(col_widths):
                for row in table.rows:
                    row.cells[i].width = Cm(width)

        self.doc.add_paragraph()  # Odstęp po tabeli

    def add_code_block(self, code: str, language: str = "") -> None:
        """Dodaje blok kodu."""
        if not self.doc:
            return

        p = self.doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def replace_placeholders(self, text: str) -> str:
        """Zamienia placeholdery na wartości z konfiguracji."""
        replacements = {
            "TODO": "",
            "[Imię Nazwisko]": self.config.get("opiekun", ""),
        }

        for placeholder, value in replacements.items():
            text = text.replace(placeholder, value)

        return text

    def save(self, path: Path) -> None:
        """Zapisuje dokument."""
        if self.doc:
            self.doc.save(str(path))
            print(f"✅ Zapisano: {path}")


class DPPGenerator(WordDocumentGenerator):
    """Generator dokumentu DPP (Dokumentacja Procesu Projektowania)."""

    def generate(self) -> Document:
        """Generuje dokument DPP."""
        self.create_document()

        # Nagłówek
        self.add_title("Dokumentacja Procesu Projektowania (DPP)")
        self.add_paragraph(f"Projekt: {self.config['temat']}", bold=True)
        self.add_paragraph(f"Rok akademicki: {self.config['rok_akademicki']}")
        self.add_paragraph("")

        # 1. Informacja o projekcie
        self.add_title("1. Informacja o projekcie", level=2)

        self.add_title("1.1 Temat projektu", level=3)
        self.add_paragraph(self.config["temat"])

        self.add_title("1.2 Cel projektu", level=3)
        self.add_paragraph(
            "Stworzenie datasetu w formacie COCO zawierającego:\n"
            "• Bounding boxes psów na obrazach/klatkach wideo\n"
            "• Klasyfikację ras psów\n"
            "• Punkty kluczowe twarzy (facial keypoints)\n"
            "• Etykiety emocji według systemu DogFACS"
        )

        self.add_title("1.3 Zespół projektowy", level=3)
        team_data = [["Lp.", "Imię i nazwisko", "Nr albumu", "Rola"]]
        for i, member in enumerate(self.config["zespol"], 1):
            team_data.append([str(i), member["imie"], member["album"], member["rola"]])
        self.add_table(team_data)
        self.add_paragraph(f"Opiekun projektu: {self.config['opiekun']}", bold=True)

        # 2. Podział zadań
        self.add_title("2. Podział zadań i ról", level=2)

        roles_data = [["Osoba", "Odpowiedzialność"]]
        role_assignments = [
            (self.config["zespol"][0]["imie"], "Architektura systemu, pipeline inference, koordynacja"),
            (self.config["zespol"][1]["imie"], "Modele AI (YOLOv8, EfficientNet), trening"),
            (self.config["zespol"][2]["imie"], "Zbieranie danych, batch processing, COCO format"),
            (self.config["zespol"][3]["imie"], "Weryfikacja manualna, testy, dokumentacja"),
        ]
        for person, role in role_assignments:
            roles_data.append([person, role])
        self.add_table(roles_data)

        # 3. Specyfikacja wymagań
        self.add_title("3. Specyfikacja wymagań", level=2)

        self.add_title("3.1 Wymagania funkcjonalne", level=3)
        func_req = [
            ["ID", "Wymaganie", "Priorytet"],
            ["WF01", "System wykrywa psy na obrazach/wideo", "Wysoki"],
            ["WF02", "System klasyfikuje rasę psa (120 klas)", "Wysoki"],
            ["WF03", "System wykrywa 46 punktów kluczowych twarzy", "Wysoki"],
            ["WF04", "System klasyfikuje emocje psa (6 kategorii)", "Wysoki"],
            ["WF05", "System eksportuje dane w formacie COCO", "Wysoki"],
            ["WF06", "Aplikacja demo umożliwia upload obrazów/wideo", "Średni"],
            ["WF07", "System wspiera batch processing", "Średni"],
        ]
        self.add_table(func_req)

        self.add_title("3.2 Wymagania niefunkcjonalne", level=3)
        nonfunc_req = [
            ["ID", "Wymaganie", "Wartość docelowa"],
            ["WN01", "Dokładność detekcji psów (mAP@0.5)", "> 0.85"],
            ["WN02", "Dokładność klasyfikacji ras (Top-5)", "> 0.90"],
            ["WN03", "Błąd keypoints (PCK@0.2)", "< 0.15"],
            ["WN04", "Czas inference na GPU", "< 500ms/obraz"],
            ["WN05", "Format wyjściowy", "COCO JSON"],
        ]
        self.add_table(nonfunc_req)

        # 4. Harmonogram
        self.add_title("4. Harmonogram prac", level=2)

        sprint_data = [["Sprint", "Nazwa", "Status"]]
        for sprint in self.config["sprinty"]:
            sprint_data.append([str(sprint["nr"]), sprint["nazwa"], sprint["status"]])
        self.add_table(sprint_data)

        # 5. Technologie
        self.add_title("5. Narzędzia i technologie", level=2)

        tech_data = [
            ["Kategoria", "Technologia"],
            ["Język programowania", "Python 3.11+"],
            ["Deep Learning", "PyTorch, Ultralytics"],
            ["Detekcja obiektów", "YOLOv8"],
            ["Klasyfikacja", "EfficientNet-B4"],
            ["Keypoints", "HRNet / SimpleBaseline"],
            ["Interfejs użytkownika", "Streamlit"],
            ["Format danych", "COCO JSON"],
            ["Wersjonowanie kodu", "Git, GitHub"],
            ["Linter / Formatter", "Ruff"],
        ]
        self.add_table(tech_data)

        self.add_paragraph(f"Repozytorium: {self.config['repo_url']}")

        # 6. Ryzyka
        self.add_title("6. Ryzyka projektu", level=2)

        risk_data = [
            ["Ryzyko", "Prawdopodobieństwo", "Wpływ", "Mitygacja"],
            ["Brak danych treningowych", "Średnie", "Wysoki", "Użycie publicznych datasetów"],
            ["Niska dokładność modeli", "Średnie", "Wysoki", "Transfer learning, fine-tuning"],
            ["Opóźnienia w harmonogramie", "Wysokie", "Średni", "Buffer czasowy, priorytyzacja"],
            ["Problemy z GPU", "Niskie", "Wysoki", "Google Colab jako backup"],
        ]
        self.add_table(risk_data)

        return self.doc


class SpecyfikacjaGenerator(WordDocumentGenerator):
    """Generator dokumentu Specyfikacja Oprogramowania."""

    def generate(self) -> Document:
        """Generuje dokument Specyfikacji."""
        self.create_document()

        # Nagłówek
        self.add_title("Specyfikacja Oprogramowania")
        self.add_paragraph(f"Projekt: {self.config['temat']}", bold=True)
        self.add_paragraph("")

        # 1. Charakterystyka funkcjonalna
        self.add_title("1. Charakterystyka funkcjonalna", level=2)

        self.add_title("1.1 Opis systemu", level=3)
        self.add_paragraph(
            "System Dog FACS Dataset to pipeline do automatycznej anotacji emocji psów "
            "wykorzystujący modele głębokiego uczenia. System przetwarza obrazy lub klatki "
            "wideo i generuje anotacje w formacie COCO."
        )

        self.add_title("1.2 Główne funkcje", level=3)

        functions = [
            ("F1: Detekcja psów", "YOLOv8-m", "Obraz RGB → Lista bounding boxes z confidence"),
            ("F2: Klasyfikacja ras", "EfficientNet-B4", "Crop psa → Top-5 ras z prawdopodobieństwami"),
            ("F3: Detekcja keypoints", "SimpleBaseline (ResNet-50)", "Crop psa → 46 punktów kluczowych twarzy"),
            ("F4: Klasyfikacja emocji", "EfficientNet-B0", "Crop psa → 6 kategorii emocji"),
            ("F5: Pipeline inference", "-", "Orkiestracja wszystkich modeli, eksport COCO"),
            ("F6: Aplikacja demo", "Streamlit", "Upload obrazów/wideo, wizualizacja wyników"),
        ]

        func_table = [["Funkcja", "Model/Technologia", "Opis"]]
        for func in functions:
            func_table.append(list(func))
        self.add_table(func_table)

        # 2. Interfejs
        self.add_title("2. Opis interfejsu", level=2)

        self.add_title("2.1 Interfejs programistyczny (API)", level=3)
        self.add_code_block(
            """from packages.pipeline import InferencePipeline

# Inicjalizacja
pipeline = InferencePipeline(device="cuda", confidence_threshold=0.5)

# Inference na obrazie
results = pipeline.process_image("image.jpg")

# Eksport COCO
pipeline.export_coco(results, "output.json")"""
        )

        self.add_title("2.2 Interfejs użytkownika", level=3)
        ui_table = [
            ["Ekran", "Opis"],
            ["Upload", "Wybór pliku obrazu lub wideo"],
            ["Processing", "Pasek postępu, podgląd przetwarzania"],
            ["Results", "Wizualizacja bbox, keypoints, emocji"],
            ["Export", "Pobranie anotacji w formacie COCO JSON"],
        ]
        self.add_table(ui_table)

        # 3. Architektura
        self.add_title("3. Opis oprogramowania", level=2)

        self.add_title("3.1 Architektura systemu", level=3)
        self.add_paragraph(
            "System składa się z trzech warstw:\n\n"
            "1. Warstwa prezentacji (Streamlit App)\n"
            "   - Interfejs użytkownika do uploadu i wizualizacji\n\n"
            "2. Warstwa logiki (Pipeline Layer)\n"
            "   - BBoxModel: detekcja psów (YOLOv8)\n"
            "   - BreedModel: klasyfikacja ras (EfficientNet-B4)\n"
            "   - KeypointModel: punkty kluczowe (SimpleBaseline)\n"
            "   - EmotionModel: klasyfikacja emocji (EfficientNet-B0)\n\n"
            "3. Warstwa danych (Data Layer)\n"
            "   - COCOReader: odczyt datasetu\n"
            "   - COCOWriter: zapis anotacji"
        )

        self.add_title("3.2 Struktura katalogów", level=3)
        self.add_code_block(
            """group-project/
├── apps/
│   ├── demo/              # Aplikacja Streamlit
│   └── verification/      # Narzędzie weryfikacji
├── packages/
│   ├── models/            # Modele AI (bbox, breed, keypoints, emotion)
│   ├── pipeline/          # Orkiestracja inference
│   └── data/              # COCO utilities
├── scripts/
│   ├── annotation/        # Batch processing
│   ├── download/          # Pobieranie danych
│   ├── training/          # Trening modeli
│   └── verification/      # Weryfikacja anotacji
├── notebooks/             # Analiza statystyk
└── docs/                  # Dokumentacja"""
        )

        # 4. Wyniki
        self.add_title("4. Wyniki", level=2)

        self.add_title("4.1 Metryki modeli", level=3)
        metrics_table = [
            ["Model", "Metryka", "Wartość docelowa", "Wartość osiągnięta"],
            ["Detector (YOLOv8)", "mAP@0.5", "> 0.85", self.config["metryki"]["detector_map"]],
            ["Breed (EfficientNet)", "Top-5 Accuracy", "> 0.90", self.config["metryki"]["breed_top5"]],
            ["Keypoints", "PCK@0.2", "> 0.85", self.config["metryki"]["keypoints_pck"]],
            ["Emotion", "F1-score", "> 0.70", self.config["metryki"]["emotion_f1"]],
        ]
        self.add_table(metrics_table)

        self.add_title("4.2 Statystyki datasetu", level=3)
        dataset_table = [
            ["Statystyka", "Wartość"],
            ["Liczba obrazów", self.config["dataset"]["liczba_obrazow"]],
            ["Liczba anotacji psów", self.config["dataset"]["liczba_anotacji"]],
            ["Unikalne rasy", self.config["dataset"]["unikalne_rasy"]],
            ["Średnia keypoints per pies", self.config["dataset"]["keypoints_per_pies"]],
        ]
        self.add_table(dataset_table)

        return self.doc


class RaportRocznyGenerator(WordDocumentGenerator):
    """Generator Raportu Rocznego."""

    def generate(self) -> Document:
        """Generuje Raport Roczny."""
        self.create_document()

        # Nagłówek
        self.add_title("Raport Roczny")
        self.add_paragraph(f"Temat: {self.config['temat']}", bold=True)
        self.add_paragraph(f"Rok akademicki: {self.config['rok_akademicki']}")
        self.add_paragraph(f"Semestr: {self.config['semestr']}")
        self.add_paragraph("")

        # 1. Wykonawcy
        self.add_title("1. Wykonawcy", level=2)
        team_data = [["Lp.", "Imię i nazwisko", "Nr albumu", "Rola"]]
        for i, member in enumerate(self.config["zespol"], 1):
            team_data.append([str(i), member["imie"], member["album"], member["rola"]])
        self.add_table(team_data)
        self.add_paragraph(f"Opiekun projektu: {self.config['opiekun']}")

        # 2. Główne zadania
        self.add_title("2. Główne zadania", level=2)

        self.add_title("2.1 Zrealizowane sprinty", level=3)
        sprint_data = [["Sprint", "Nazwa", "Status"]]
        for sprint in self.config["sprinty"]:
            sprint_data.append([str(sprint["nr"]), sprint["nazwa"], sprint["status"]])
        self.add_table(sprint_data)

        self.add_title("2.2 Podział pracy", level=3)
        work_data = [
            ["Osoba", "Główne zadania", "Wkład"],
            [self.config["zespol"][0]["imie"], "Architektura, pipeline, koordynacja", "25%"],
            [self.config["zespol"][1]["imie"], "Modele AI, trening", "25%"],
            [self.config["zespol"][2]["imie"], "Zbieranie danych, batch processing", "25%"],
            [self.config["zespol"][3]["imie"], "Weryfikacja, testy, dokumentacja", "25%"],
        ]
        self.add_table(work_data)

        # 3. Osiągnięte wyniki
        self.add_title("3. Osiągnięte wyniki", level=2)

        self.add_title("3.1 Produkt końcowy", level=3)
        self.add_paragraph(
            "• Dataset COCO z anotacjami emocji psów\n"
            "• 4 wytrenowane modele AI\n"
            "• Aplikacja demonstracyjna (Streamlit)\n"
            "• Kompletna dokumentacja techniczna"
        )

        self.add_title("3.2 Metryki techniczne", level=3)
        metrics_table = [
            ["Model", "Metryka", "Wartość"],
            ["Detector (YOLOv8)", "mAP@0.5", self.config["metryki"]["detector_map"]],
            ["Breed (EfficientNet)", "Top-5 Accuracy", self.config["metryki"]["breed_top5"]],
            ["Keypoints", "PCK@0.2", self.config["metryki"]["keypoints_pck"]],
            ["Emotion", "F1-score", self.config["metryki"]["emotion_f1"]],
        ]
        self.add_table(metrics_table)

        # 4. Dokumentacja
        self.add_title("4. Dokumentacja", level=2)
        doc_table = [
            ["Dokument", "Lokalizacja"],
            ["DPP", "docs/deliverables/DPP.md"],
            ["Specyfikacja", "docs/deliverables/Specyfikacja.md"],
            ["Repozytorium", self.config["repo_url"]],
        ]
        self.add_table(doc_table)

        # 5. Wnioski
        self.add_title("5. Wnioski", level=2)

        self.add_title("5.1 Co się udało", level=3)
        self.add_paragraph(
            "• Zaimplementowano kompletny pipeline AI\n"
            "• Stworzono modułową architekturę\n"
            "• Przygotowano narzędzia do weryfikacji\n"
            "• Opracowano dokumentację"
        )

        self.add_title("5.2 Napotkane problemy", level=3)
        self.add_paragraph(
            "• Ograniczona dostępność danych treningowych dla emocji\n"
            "• Subiektywność etykiet emocji\n"
            "• Czas treningu modeli na ograniczonym GPU"
        )

        self.add_title("5.3 Rekomendacje na przyszłość", level=3)
        self.add_paragraph(
            "• Rozszerzenie o pełny system DogFACS\n"
            "• Active learning dla trudnych przypadków\n"
            "• Integracja z aplikacjami mobilnymi"
        )

        return self.doc


class PlakatGenerator(WordDocumentGenerator):
    """Generator Plakatu."""

    def generate(self) -> Document:
        """Generuje Plakat."""
        self.create_document()

        # Nagłówek
        self.add_paragraph("POLITECHNIKA GDAŃSKA", bold=True)
        self.add_paragraph("Wydział Elektroniki, Telekomunikacji i Informatyki")
        self.add_paragraph(f"Katedra: {self.config['katedra']}")
        self.add_paragraph("")

        # Tytuł
        self.add_title("Dog FACS Dataset")
        self.add_paragraph(
            "Pipeline do automatycznej anotacji emocji psów z wykorzystaniem AI",
            italic=True,
        )
        self.add_paragraph("")

        # Zespół
        self.add_title("Zespół", level=2)
        for member in self.config["zespol"]:
            self.add_paragraph(f"• {member['imie']} - {member['rola']}")
        self.add_paragraph(f"Opiekun: {self.config['opiekun']}")
        self.add_paragraph("")

        # Cel
        self.add_title("Cel projektu", level=2)
        self.add_paragraph(
            "Stworzenie datasetu w formacie COCO zawierającego anotacje emocji psów "
            "wykrywanych automatycznie przez modele głębokiego uczenia."
        )

        # Architektura
        self.add_title("Architektura", level=2)
        arch_table = [
            ["Etap", "Model", "Zadanie"],
            ["1", "YOLOv8", "Detekcja psów"],
            ["2", "EfficientNet-B4", "Klasyfikacja ras"],
            ["3", "SimpleBaseline", "46 keypoints twarzy"],
            ["4", "EfficientNet-B0", "6 kategorii emocji"],
        ]
        self.add_table(arch_table)

        # Wyniki
        self.add_title("Wyniki", level=2)
        results_table = [
            ["Metryka", "Wartość"],
            ["Obrazów w datasecie", self.config["dataset"]["liczba_obrazow"]],
            ["Anotacji psów", self.config["dataset"]["liczba_anotacji"]],
            ["Dokładność detekcji", self.config["metryki"]["detector_map"]],
        ]
        self.add_table(results_table)

        # Technologie
        self.add_title("Technologie", level=2)
        self.add_paragraph("Python • PyTorch • Streamlit • YOLOv8 • COCO JSON")

        # Kontakt
        self.add_paragraph("")
        self.add_paragraph(f"Repozytorium: {self.config['repo_url']}")
        self.add_paragraph("Projekt grupowy, semestr zimowy 2025/2026", italic=True)

        return self.doc


# Mapa generatorów
GENERATORS = {
    "DPP": DPPGenerator,
    "Specyfikacja": SpecyfikacjaGenerator,
    "Raport-roczny": RaportRocznyGenerator,
    "Plakat": PlakatGenerator,
}


def list_documents() -> None:
    """Wyświetla listę dostępnych dokumentów."""
    print("\nDostępne dokumenty do wygenerowania:\n")
    for name in GENERATORS:
        print(f"  • {name}")
    print(f"\nUżycie: python {sys.argv[0]} --document <nazwa>")
    print(f"        python {sys.argv[0]} --all")


def generate_document(name: str, output_dir: Path) -> Optional[Path]:
    """Generuje pojedynczy dokument."""
    if name not in GENERATORS:
        print(f"❌ Nieznany dokument: {name}")
        print(f"   Dostępne: {', '.join(GENERATORS.keys())}")
        return None

    generator_class = GENERATORS[name]
    generator = generator_class(PROJECT_CONFIG)
    generator.generate()

    output_path = output_dir / f"{name}.docx"
    generator.save(output_path)
    return output_path


def main() -> None:
    """Główna funkcja."""
    parser = argparse.ArgumentParser(
        description="Generator dokumentów Word z szablonów projektu",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Przykłady:
  %(prog)s --list                    # Lista dostępnych dokumentów
  %(prog)s --all                     # Generuj wszystkie dokumenty
  %(prog)s --document DPP            # Generuj tylko DPP
  %(prog)s --document Specyfikacja --output ./output/
        """,
    )

    parser.add_argument(
        "--document",
        "-d",
        type=str,
        help="Nazwa dokumentu do wygenerowania",
    )
    parser.add_argument(
        "--all",
        "-a",
        action="store_true",
        help="Generuj wszystkie dokumenty",
    )
    parser.add_argument(
        "--list",
        "-l",
        action="store_true",
        help="Wyświetl listę dostępnych dokumentów",
    )
    parser.add_argument(
        "--output",
        "-o",
        type=Path,
        default=Path("docs/generated"),
        help="Katalog wyjściowy (domyślnie: docs/generated)",
    )

    args = parser.parse_args()

    if args.list:
        list_documents()
        return

    if not args.document and not args.all:
        parser.print_help()
        return

    # Utwórz katalog wyjściowy
    args.output.mkdir(parents=True, exist_ok=True)

    if args.all:
        print(f"\n📄 Generowanie wszystkich dokumentów do {args.output}/\n")
        for name in GENERATORS:
            generate_document(name, args.output)
        print(f"\n✅ Wygenerowano {len(GENERATORS)} dokumentów")
    else:
        generate_document(args.document, args.output)


if __name__ == "__main__":
    main()
