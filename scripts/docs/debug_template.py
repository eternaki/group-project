#!/usr/bin/env python3
"""Diagnostyka struktury szablonu."""

import sys
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document


def main():
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "template" / "PG_WETI_DTP_wer. 1.00.docx"

    doc = Document(docx_path)

    print("=== STRUKTURA TABELI ===")

    if doc.tables:
        table = doc.tables[0]

        # Показать только первые 5 строк для диагностики
        for row_idx, row in enumerate(table.rows[:5]):
            print(f"\n--- ROW {row_idx} ---")
            for cell_idx, cell in enumerate(row.cells[:3]):
                print(f"\n  CELL {cell_idx}:")
                print(f"    Full text: {cell.text[:80]}...")
                print(f"    Paragraphs: {len(cell.paragraphs)}")

                for para_idx, para in enumerate(cell.paragraphs):
                    print(f"\n    Para {para_idx}: {len(para.runs)} runs")
                    for run_idx, run in enumerate(para.runs):
                        text = run.text.replace('\n', '\\n')[:50]
                        has_placeholder = '{' in run.text and '}' in run.text
                        print(f"      Run {run_idx}: '{text}' [placeholder: {has_placeholder}]")


if __name__ == "__main__":
    main()
