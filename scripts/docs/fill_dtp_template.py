#!/usr/bin/env python3
"""
Skrypt do wypełnienia szablonu DTP danymi projektu.
Zachowuje oryginalny styl dokumentu.
"""

import re
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color: str):
    """Ustawia kolor tła komórki tabeli."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def parse_markdown_table(text: str) -> list[list[str]] | None:
    """Parsuje tabelę markdown i zwraca listę wierszy."""
    lines = text.strip().split('\n')
    rows = []
    for line in lines:
        if line.strip().startswith('|') and '---' not in line:
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)
    return rows if len(rows) > 1 else None


def add_word_table(doc, paragraph, rows: list[list[str]], title: str = None):
    """Dodaje tabelę Word po wskazanym paragrafie."""
    if not rows:
        return

    # Dodaj tytuł tabeli jeśli podany
    if title:
        title_para = doc.add_paragraph()
        title_para.add_run(title).bold = True
        paragraph._element.addnext(title_para._element)
        paragraph = title_para

    # Utwórz tabelę
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                cell.text = cell_text.strip()

                # Nagłówek tabeli - pogrubiony z szarym tłem
                if i == 0:
                    set_cell_shading(cell, "D9D9D9")
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                        para.paragraph_format.space_after = Pt(0)
                else:
                    for para in cell.paragraphs:
                        para.paragraph_format.space_after = Pt(0)

    # Wstaw tabelę po paragrafie
    paragraph._element.addnext(table._tbl)

    return table


def insert_content_with_tables(doc, paragraph, content: str):
    """Wstawia zawartość z automatyczną konwersją tabel markdown na tabele Word."""
    lines = content.strip().split('\n')
    current_para = paragraph
    table_lines = []
    in_table = False
    table_title = None

    i = 0
    while i < len(lines):
        line = lines[i]

        # Wykryj początek tabeli
        if line.strip().startswith('|') and not in_table:
            in_table = True
            table_lines = [line]
            # Sprawdź czy poprzednia linia to tytuł tabeli
            if i > 0 and lines[i-1].strip().startswith('Tabela'):
                table_title = lines[i-1].strip()
            i += 1
            continue

        # Kontynuuj zbieranie linii tabeli
        if in_table:
            if line.strip().startswith('|') or line.strip().startswith('|-'):
                table_lines.append(line)
                i += 1
                continue
            else:
                # Koniec tabeli - utwórz tabelę Word
                in_table = False
                table_text = '\n'.join(table_lines)
                rows = parse_markdown_table(table_text)
                if rows:
                    tbl = add_word_table(doc, current_para, rows, table_title)
                    if tbl:
                        # Znajdź paragraf po tabeli
                        new_para = doc.add_paragraph()
                        tbl.addnext(new_para._element)
                        current_para = new_para
                table_lines = []
                table_title = None
                # Nie zwiększamy i - przetworzymy tę linię ponownie

        # Zwykły tekst (nie tabela)
        if not in_table:
            # Pomiń linie z tytułem tabeli (już obsłużone)
            if line.strip().startswith('Tabela') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
                i += 1
                continue

            # Dodaj paragraf z tekstem
            if line.strip():
                new_para = doc.add_paragraph(line)
                current_para._element.addnext(new_para._element)
                current_para = new_para

            i += 1

    # Obsłuż tabelę na końcu
    if in_table and table_lines:
        table_text = '\n'.join(table_lines)
        rows = parse_markdown_table(table_text)
        if rows:
            add_word_table(doc, current_para, rows, table_title)

    return current_para


# ============================================================================
# DANE PROJEKTU
# ============================================================================

PROJECT_DATA = {
    # Metadata
    "nazwa_projektu": "Dog FACS Dataset",
    "akronim": "DogFACS",
    "zleceniodawca": "dr hab. inż. Michał Czubenko",
    "numer_zlecenia": "PG/WETI/2025/01",
    "kierownik_projektu": "Danylo Lohachov",
    "opiekun_projektu": "dr hab. inż. Michał Czubenko",
    "nazwa_dokumentu": "Dokumentacja Techniczna Projektu",
    "akronim_dokumentu": "DTP",
    "wersja": "2.00",
    "odpowiedzialny": "Danylo Lohachov",
    "data_pierwsza": "28.01.2025",
    "data_aktualizacji": datetime.now().strftime("%d.%m.%Y"),
    "kierunek": "Informatyka",
    "semestr": "I",

    # Historia zmian
    "historia": [
        ("1.00", "Wersja wstępna", "całość", "Lohachov, Danylo", "28.01.2025"),
        ("2.00", "Aktualizacja na podstawie gałęzi feature/dogfacs-rule-based", "całość", "Lohachov, Danylo", "28.01.2025"),
    ],

    # Zespół
    "zespol": [
        ("Danylo Lohachov", "196610", "Kierownik projektu / Dokumentacja / QA / Frontend"),
        ("Anton Shkrebela", "196637", "AI/ML Specialist (Keypoints & DogFACS)"),
        ("Danylo Zherzdiev", "196765", "Backend (BBox & Breed models, Pipeline, COCO)"),
        ("Mariia Volkova", "196660", "Data Engineer (Data collection & Manual verification)"),
    ],
}


# ============================================================================
# ZAWARTOŚĆ SEKCJI
# ============================================================================

SECTION_1_2_ZAKRES = """Dokument obejmuje:
• Opis architektury systemu Dog FACS Dataset
• Specyfikację funkcjonalną pipeline'u AI
• Parametry techniczne modeli głębokiego uczenia (YOLOv8m, EfficientNet-B4, SimpleBaseline)
• Specyfikację systemu DogFACS Rule Engine do klasyfikacji emocji
• Schematy blokowe przetwarzania danych
• Opis interfejsu programistycznego (API) i REST API
• Format wyjściowy COCO z rozszerzeniami DogFACS
• Instrukcję instalacji i uruchomienia

Dokument nie obejmuje:
• Szczegółowej dokumentacji kodu źródłowego (dostępna w repozytorium)
• Instrukcji trenowania modeli (dostępna w docs/plans/)
• Danych treningowych (gitignored, dostępne lokalnie)

Dokumenty powiązane:
• DPP - Dokumentacja Procesu Projektowania
• docs/plans/2025-01-03-architecture.md - Architektura systemu
• README.md - Instrukcja szybkiego startu"""

SECTION_1_3_ODBIORCY = """Dokument jest przeznaczony dla:
• Zleceniodawca: Katedra Systemów Decyzyjnych i Robotyki, Politechnika Gdańska
• Opiekun projektu: dr hab. inż. Michał Czubenko
• Komisja oceniająca projekt grupowy WETI

Członkowie zespołu projektowego:
• Danylo Lohachov (196610) - Kierownik projektu
• Anton Shkrebela (196637) - AI/ML Specialist
• Danylo Zherzdiev (196765) - Backend Developer
• Mariia Volkova (196660) - Data Engineer

Osoby zainteresowane rozwojem lub użytkowaniem systemu."""

SECTION_1_4_TERMINOLOGIA = """Tabela 1.1. Terminologia

| Termin | Definicja |
|--------|-----------|
| DogFACS | Dog Facial Action Coding System - naukowy system kodowania mimiki psów |
| COCO | Common Objects in Context - format anotacji obrazów |
| Keypoints | Punkty kluczowe - 20 charakterystycznych punktów anatomicznych twarzy psa |
| BBox | Bounding Box - prostokąt obejmujący obiekt na obrazie |
| Action Unit (AU) | Jednostka akcji mimicznej - obiektywna miara ruchu mięśni twarzy |
| Delta AU | Różnica wartości AU względem neutralnej klatki bazowej |
| Neutral Frame | Klatka referencyjna z neutralnym wyrazem twarzy psa |
| Peak Frame | Klatka z maksymalną ekspresją emocji (wysoki TFM) |
| TFM | Total Facial Movement - suma wszystkich aktywacji AU |
| Rule-based | Klasyfikacja oparta na regułach naukowych, bez uczenia maszynowego |
| Poselet | Kombinacja AU charakterystyczna dla określonej emocji |
| YOLOv8 | You Only Look Once v8 - architektura detekcji obiektów |
| EfficientNet | Efficient Neural Network - architektura klasyfikacji obrazów |
| SimpleBaseline | Architektura detekcji keypoints (ResNet + Deconv) |"""


SECTION_2_CONTENT = """2.1	Opis produktu

System Dog FACS Dataset to pipeline do automatycznej anotacji emocji psów wykorzystujący:
• 3 modele deep learning (detekcja, klasyfikacja ras, keypoints)
• Rule-based DogFACS engine do klasyfikacji emocji (BEZ uczenia maszynowego)

System przetwarza wideo i generuje anotacje w formacie COCO zawierające:
• Bounding boxes - prostokąty obejmujące wykryte psy
• Klasyfikację ras - identyfikacja rasy psa (50+ ras)
• Punkty kluczowe twarzy - 20 punktów anatomicznych
• Action Units - 12 jednostek akcji mimicznej DogFACS
• Etykiety emocji - 6 klas (happy, sad, angry, fearful, relaxed, neutral)

2.2	Architektura systemu

2.2.1	Schemat blokowy systemu

Video Input → Frame Extraction → Neutral Frame Detection
                                        ↓
                              AI Pipeline:
                              ├── BBox (YOLOv8m)
                              ├── Breed (EfficientNet-B4)
                              └── Keypoints (SimpleBaseline)
                                        ↓
                              Delta Action Units Extractor
                              (neutral vs target frame)
                                        ↓
                              DogFACS Rule Engine (NO ML)
                              (AU → Poselet → Emotion)
                                        ↓
                              Peak Frame Selector (TFM)
                                        ↓
                              COCO Exporter → Dataset

2.2.2	Architektura warstw

┌─────────────────────────────────────────────────────────────────┐
│                    APPLICATION LAYER                             │
│  Streamlit Demo | React + FastAPI Webapp | Verification Tool    │
└─────────────────────────────────────────────────────────────────┘
                              ↓
┌─────────────────────────────────────────────────────────────────┐
│                     PIPELINE LAYER                               │
│  inference.py | neutral_frame.py | peak_selector.py             │
│  video.py | temporal_processor.py                               │
└─────────────────────────────────────────────────────────────────┘
                              ↓
┌─────────────────────────────────────────────────────────────────┐
│                      MODELS LAYER                                │
│  bbox.py (YOLOv8m) | breed.py (EffNetB4) | keypoints.py         │
│  emotion.py (Rule-based) | action_units.py | head_pose.py       │
└─────────────────────────────────────────────────────────────────┘
                              ↓
┌─────────────────────────────────────────────────────────────────┐
│                       DATA LAYER                                 │
│  coco.py (COCO format) | schemas.py (data classes)              │
└─────────────────────────────────────────────────────────────────┘

2.3	Specyfikacja modeli AI

Tabela 2.1. Specyfikacja modeli

| Model | Plik wag | Rozmiar | Architektura | Wejście | Wyjście |
|-------|----------|---------|--------------|---------|---------|
| BBox | yolov8m.pt | 52.1 MB | YOLOv8m | 640×640 px | Lista bbox + confidence |
| Breed | breed.pt | 71.8 MB | EfficientNet-B4 | 224×224 px | Top-5 ras |
| Keypoints | keypoints_dogflw.pt | 102.1 MB | SimpleBaseline (ResNet34) | 256×256 px | 20 punktów (x, y, vis) |
| Emotion | BRAK (rule-based) | - | DogFACS Rule Engine | 12 Delta AU | 6 klas emocji |

2.3.1	Lista 20 keypoints projektu

| ID | Nazwa | Opis |
|----|-------|------|
| 0 | left_eye | Środek lewego oka |
| 1 | right_eye | Środek prawego oka |
| 2 | nose | Czubek nosa |
| 3 | left_ear_base | Podstawa lewego ucha |
| 4 | right_ear_base | Podstawa prawego ucha |
| 5 | left_ear_tip | Czubek lewego ucha |
| 6 | right_ear_tip | Czubek prawego ucha |
| 7 | left_mouth_corner | Lewy kącik ust |
| 8 | right_mouth_corner | Prawy kącik ust |
| 9 | upper_lip | Środek górnej wargi |
| 10 | lower_lip | Środek dolnej wargi |
| 11 | chin | Podbródek |
| 12 | left_cheek | Lewy policzek |
| 13 | right_cheek | Prawy policzek |
| 14 | forehead | Środek czoła |
| 15 | left_eyebrow | Lewa brew |
| 16 | right_eyebrow | Prawa brew |
| 17 | muzzle_top | Góra pyska |
| 18 | muzzle_left | Lewa strona pyska |
| 19 | muzzle_right | Prawa strona pyska |

2.4	System DogFACS Action Units

2.4.1	Lista oficjalnych Action Units

System implementuje 12 oficjalnych kodów DogFACS:

| Kod AU | Nazwa | Opis |
|--------|-------|------|
| AU101 | Inner Brow Raiser | Podniesienie wewnętrznej brwi |
| AU102 | Outer Brow Raiser | Podniesienie zewnętrznej brwi |
| AU12 | Lip Corner Puller | Pociągnięcie kącików ust (uśmiech) |
| AU115 | Upper Eyelid Raiser | Podniesienie górnej powieki |
| AU116 | Lower Eyelid Raiser | Zmrużenie dolnej powieki |
| AU117 | Closure of Eyelids | Zamknięcie oczu (mruganie) |
| AU121 | Eye Widener | Rozszerzenie oczu |
| EAD102 | Ears Forward | Uszy do przodu |
| EAD103 | Ears Flattener | Uszy spłaszczone/do tyłu |
| AD19 | Tongue Show | Pokazanie języka |
| AD37 | Nose Lick | Lizanie nosa |
| AU26 | Jaw Drop | Opadnięcie szczęki |

2.4.2	Delta AU Extraction

Zasada działania:
Delta_AU = (distance_target / distance_neutral) - 1.0

Gdzie:
• distance_target - pomiar geometryczny na klatce docelowej
• distance_neutral - pomiar na neutralnej klatce bazowej
• Wynik > 0 oznacza aktywację AU, < 0 oznacza deaktywację

2.4.3	Reguły klasyfikacji emocji (Poselets)

HAPPY (priorytet 100):
Wymagane: AU12 ≥ 1.20, EAD102 ≥ 1.10
Inhibitory: EAD103 < 1.10, AU26 < 1.25
Opcjonalne: AU101 ≥ 1.10 (bonus)

ANGRY (priorytet 95):
Wymagane: AU26 ≥ 1.25, AU12 ≥ 1.15
Inhibitory: EAD102 < 1.10

FEARFUL (priorytet 90):
Wymagane: EAD103 ≥ 1.15, AU101 ≥ 1.10
Inhibitory: AU26 < 1.20

SAD (priorytet 85):
Wymagane: EAD103 ≥ 1.10
Inhibitory: AU26 < 1.15, AU12 < 1.10

RELAXED (priorytet 70):
Wymagane: brak silnych aktywacji
Wszystkie AU < 1.10

NEUTRAL (priorytet 50):
Fallback - zawsze dopasowuje jeśli żadna inna reguła nie pasuje

2.5	Interfejs programistyczny (API)

2.5.1	Pipeline podstawowy (Python)

from packages.pipeline import InferencePipeline, PipelineConfig

config = PipelineConfig(
    device="cuda",
    confidence_threshold=0.3,
    use_rule_based_emotion=True
)

pipeline = InferencePipeline(config)
pipeline.load()
result = pipeline.process_frame(image)

2.5.2	REST API (FastAPI)

| Endpoint | Metoda | Opis |
|----------|--------|------|
| /api/health | GET | Health check |
| /api/process_video | POST | Przetwarza wideo, zwraca peak frames |
| /api/export_coco | POST | Eksportuje dataset do formatu COCO |
| /static/frames/* | GET | Dostęp do zapisanych klatek |

2.6	Format wyjściowy COCO (rozszerzony)

{
  "info": { "description": "Dog FACS Dataset", "version": "1.0" },
  "images": [{
    "id": 1,
    "file_name": "video001_frame_00150.jpg",
    "source_video": "video001.mp4",
    "frame_number": 150,
    "is_neutral_frame": false,
    "tfm_score": 0.342
  }],
  "annotations": [{
    "id": 1,
    "image_id": 1,
    "bbox": [100, 150, 300, 400],
    "keypoints": [120.5, 180.3, 2, ...],
    "breed": "golden_retriever",
    "emotion": "happy",
    "emotion_rule_applied": "happy_priority_100",
    "action_units": { "AU101": 1.15, "AU12": 1.32, "EAD102": 1.18 }
  }]
}

2.7	Wymagania systemowe

Tabela 2.2. Wymagania systemowe

| Komponent | Minimalne | Zalecane |
|-----------|-----------|----------|
| System operacyjny | Windows 10/11, Linux, macOS | - |
| Python | 3.10+ | 3.11+ |
| RAM | 8 GB | 16 GB |
| Dysk | 15 GB | 50 GB |
| GPU | - | NVIDIA CUDA 11.8+ |
| VRAM | - | 6 GB+ |
| FFmpeg | Wymagany | - |

2.8	Stack technologiczny

| Kategoria | Technologia | Wersja |
|-----------|-------------|--------|
| Runtime | Python | 3.10+ |
| ML Framework | PyTorch | 2.0+ |
| Detection | Ultralytics (YOLOv8) | 8.0+ |
| Classification | timm (EfficientNet) | 0.9+ |
| Frontend Demo | Streamlit | 1.28+ |
| Frontend Webapp | React + Vite | 18.0+ |
| Backend API | FastAPI | 0.100+ |
| Linting | Ruff | 0.1+ |

2.9	Repozytorium

| Pole | Wartość |
|------|---------|
| URL | https://github.com/eternaki/group-project |
| Gałąź główna | main |
| Gałąź feature | feature/dogfacs-rule-based |
| Struktura gałęzi | main → develop → sprint-X |
| Model storage | Git LFS |"""


ATTACHMENTS = [
    ("Schemat blokowy systemu", "DTP_schemat_blokowy.png"),
    ("Architektura warstw", "DTP_architektura.png"),
    ("Specyfikacja modeli AI", "DTP_modele.xlsx"),
    ("Lista Action Units DogFACS", "DTP_action_units.xlsx"),
    ("Format COCO rozszerzony", "DTP_coco_schema.json"),
    ("Dokumentacja API (Swagger)", "http://localhost:8000/docs"),
]


# ============================================================================
# FUNKCJE POMOCNICZE
# ============================================================================

def replace_text_in_paragraph(paragraph, old_text: str, new_text: str):
    """Zamienia tekst w paragrafie zachowując formatowanie."""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)


def replace_text_in_cell(cell, old_text: str, new_text: str):
    """Zamienia tekst w komórce tabeli."""
    for paragraph in cell.paragraphs:
        replace_text_in_paragraph(paragraph, old_text, new_text)


def find_paragraph_by_text(doc, search_text: str):
    """Znajduje paragraf zawierający podany tekst."""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i, para
    return None, None


def clear_paragraph(paragraph):
    """Czyści zawartość paragrafu."""
    for run in paragraph.runs:
        run.text = ""


def add_content_after_paragraph(doc, para_index: int, content: str):
    """Dodaje zawartość po wskazanym paragrafie."""
    # Usuń komentarz w nawiasach klamrowych
    para = doc.paragraphs[para_index]
    if para.text.strip().startswith('{'):
        clear_paragraph(para)

    # Dodaj nową zawartość
    lines = content.strip().split('\n')
    insert_index = para_index + 1

    for line in lines:
        if line.strip():
            # Wstawiamy nowy paragraf
            new_para = doc.paragraphs[para_index]._element
            new_p = OxmlElement('w:p')
            new_para.addnext(new_p)

            # Dodaj tekst
            from docx.text.paragraph import Paragraph
            p = Paragraph(new_p, doc.paragraphs[para_index]._parent)
            p.add_run(line)


def fill_header_table(table, data: dict):
    """Wypełnia tabelę nagłówkową danymi projektu."""
    # Mapowanie pozycji w tabeli na dane
    replacements = {
        "{nazwa i akronim projektu": f"{data['nazwa_projektu']} ({data['akronim']})",
        "{nazwa zleceniodawcy}": data['zleceniodawca'],
        "{numer zlecenia}": data['numer_zlecenia'],
        "{nazwisko, imię}": data['kierownik_projektu'],
        "{opiekun projektu}": data['opiekun_projektu'],
        "{wersja dokumentu}": data['wersja'],
        "{data sporządzenia}": data['data_pierwsza'],
        "{data aktualizacji}": data['data_aktualizacji'],
        "{kierunek studiów}": data['kierunek'],
        "{semestr}": data['semestr'],
    }

    for row in table.rows:
        for cell in row.cells:
            text_lower = cell.text.lower()
            for pattern, replacement in replacements.items():
                if pattern.lower() in text_lower:
                    # Zamień tylko placeholder
                    for para in cell.paragraphs:
                        for run in para.runs:
                            for p, r in replacements.items():
                                if p.lower() in run.text.lower():
                                    # Znajdź i zamień placeholder
                                    import re
                                    run.text = re.sub(
                                        r'\{[^}]+\}',
                                        r,
                                        run.text,
                                        count=1
                                    )


def replace_placeholder_in_paragraph(para, replacement: str):
    """Zamienia placeholder {xxx} w paragrafie, nawet jeśli jest rozdzielony na wiele runs."""
    full_text = para.text

    # Sprawdź czy jest placeholder
    if '{' not in full_text or '}' not in full_text:
        return False

    # Nie zamieniaj jeśli to {nie zmieniać}
    if 'nie zmienia' in full_text.lower():
        return False

    # Zamień placeholder w pełnym tekście
    new_text = re.sub(r'\{[^}]+\}', replacement, full_text)

    # Jeśli tekst się zmienił
    if new_text != full_text:
        # Wyczyść wszystkie runs i wstaw nowy tekst do pierwszego
        if para.runs:
            # Zachowaj pierwszy run do wstawienia tekstu
            for i, run in enumerate(para.runs):
                if i == 0:
                    run.text = new_text
                else:
                    run.text = ""
        return True
    return False


def fill_template(template_path: Path, output_path: Path):
    """Wypełnia szablon DTP danymi projektu."""
    doc = Document(template_path)

    # 1. Wypełnij tabelę nagłówkową (pierwsza tabela)
    if doc.tables:
        header_table = doc.tables[0]

        # Mapowanie kontekstu komórki na wartość do wstawienia
        # Klucz: słowo kluczowe w tekście komórki, Wartość: zamiennik
        context_map = {
            "nazwa i akronim": f"{PROJECT_DATA['nazwa_projektu']} ({PROJECT_DATA['akronim']})",
            "zleceniodawca": PROJECT_DATA['zleceniodawca'],
            "numer zlecenia": PROJECT_DATA['numer_zlecenia'],
            "kierownik projektu": PROJECT_DATA['kierownik_projektu'],
            "opiekun projektu": PROJECT_DATA['opiekun_projektu'],
            "nr wersji": PROJECT_DATA['wersja'],
            "odpowiedzialny za dokument": PROJECT_DATA['odpowiedzialny'],
            "data pierwszego": PROJECT_DATA['data_pierwsza'],
            "data ostatniej": PROJECT_DATA['data_aktualizacji'],
            "semestr realizacji": "",  # Puste - nie zamieniamy, zostaw {nie zmieniać}
        }

        # Przetwarzaj każdą komórkę tabeli
        for row in header_table.rows:
            for cell in row.cells:
                cell_text_lower = cell.text.lower()

                # Znajdź odpowiedni kontekst dla tej komórki
                replacement = None
                for keyword, value in context_map.items():
                    if keyword in cell_text_lower:
                        replacement = value
                        break

                # Jeśli znaleziono kontekst i mamy wartość do wstawienia
                if replacement is not None and replacement != "":
                    for para in cell.paragraphs:
                        replace_placeholder_in_paragraph(para, replacement)

        # Wypełnij historię zmian (wiersze 10-11 w tabeli)
        for i, (wersja, opis, rozdzial, autor, data) in enumerate(PROJECT_DATA['historia']):
            row_idx = 10 + i
            if row_idx < len(header_table.rows):
                row = header_table.rows[row_idx]
                if len(row.cells) >= 6:
                    row.cells[0].paragraphs[0].clear()
                    row.cells[0].paragraphs[0].add_run(wersja)

                    row.cells[1].paragraphs[0].clear()
                    row.cells[1].paragraphs[0].add_run(opis)

                    row.cells[2].paragraphs[0].clear()
                    row.cells[2].paragraphs[0].add_run(rozdzial)

                    # Komórki 3 i 4 są połączone dla autora
                    row.cells[3].paragraphs[0].clear()
                    row.cells[3].paragraphs[0].add_run(autor)

                    row.cells[5].paragraphs[0].clear()
                    row.cells[5].paragraphs[0].add_run(data)

    # 2. Wypełnij sekcje treści
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        # Usuń komentarze {UWAGA...} na początku dokumentu
        if text.startswith('{UWAGA') or text.startswith('{wersja dokumentu'):
            clear_paragraph(para)
            continue

        # 1.2 Zakres dokumentu
        if text.startswith('{Określenie, co wchodzi'):
            para.clear()
            # Wstaw pierwszą linię, resztę jako nowe paragrafy
            lines = SECTION_1_2_ZAKRES.strip().split('\n')
            para.add_run(lines[0] if lines else "")
            for line in lines[1:]:
                new_para = doc.add_paragraph(line)
                para._element.addnext(new_para._element)
                para = new_para

        # 1.3 Odbiorcy
        if text.startswith('{Określenie adresatów'):
            para.clear()
            lines = SECTION_1_3_ODBIORCY.strip().split('\n')
            para.add_run(lines[0] if lines else "")
            for line in lines[1:]:
                new_para = doc.add_paragraph(line)
                para._element.addnext(new_para._element)
                para = new_para

        # 1.4 Terminologia
        if text.startswith('{Wyjaśnienie używanych'):
            para.clear()
            lines = SECTION_1_4_TERMINOLOGIA.strip().split('\n')
            para.add_run(lines[0] if lines else "")
            for line in lines[1:]:
                new_para = doc.add_paragraph(line)
                para._element.addnext(new_para._element)
                para = new_para

        # 2. Dokumentacja techniczna projektu
        if text.startswith('{O zakresie dokumentacji'):
            para.clear()
            lines = SECTION_2_CONTENT.strip().split('\n')
            para.add_run(lines[0] if lines else "")
            for line in lines[1:]:
                new_para = doc.add_paragraph(line)
                para._element.addnext(new_para._element)
                para = new_para

        # 3. Załączniki
        if text.startswith('{Wszelkie dokumenty'):
            para.clear()
            para.add_run("Poniższa tabela zawiera listę załączników do dokumentu.")

    # 3. Wypełnij tabelę załączników (druga tabela)
    if len(doc.tables) > 1:
        attach_table = doc.tables[1]

        for i, (nazwa, plik) in enumerate(ATTACHMENTS):
            row_idx = i + 1  # Pomijamy nagłówek
            if row_idx < len(attach_table.rows):
                row = attach_table.rows[row_idx]
                if len(row.cells) >= 3:
                    row.cells[0].paragraphs[0].clear()
                    row.cells[0].paragraphs[0].add_run(f"{i + 1}.")

                    row.cells[1].paragraphs[0].clear()
                    row.cells[1].paragraphs[0].add_run(nazwa)

                    row.cells[2].paragraphs[0].clear()
                    row.cells[2].paragraphs[0].add_run(plik)

    # Zapisz dokument
    doc.save(output_path)
    print(f"Zapisano: {output_path}")


def main():
    """Główna funkcja."""
    project_root = Path(__file__).parent.parent.parent

    template_path = project_root / "template" / "PG_WETI_DTP_wer. 1.00.docx"
    output_path = project_root / "docs" / "deliverables" / "DTP_v3.docx"

    if not template_path.exists():
        print(f"Błąd: Nie znaleziono szablonu {template_path}")
        print("Uruchom najpierw: python scripts/docs/convert_doc.py")
        return

    fill_template(template_path, output_path)
    print(f"\nSzablon DTP wypełniony!")
    print(f"Plik wyjściowy: {output_path}")


if __name__ == "__main__":
    main()
