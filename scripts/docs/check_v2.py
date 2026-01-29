#!/usr/bin/env python3
"""Sprawdź wypełniony szablon DTP_v2."""

from pathlib import Path
from docx import Document


def main():
    project_root = Path(__file__).parent.parent.parent
    docx_path = project_root / "docs" / "deliverables" / "DTP_v2.docx"

    if not docx_path.exists():
        print(f"Błąd: Nie znaleziono pliku {docx_path}")
        return

    doc = Document(docx_path)

    print("=" * 60)
    print("WERYFIKACJA DTP_v2.docx")
    print("=" * 60)

    print("\n--- TABELA NAGŁÓWKOWA (pierwsze 8 wierszy) ---")
    if doc.tables:
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows[:8]):
            cells = [cell.text.strip()[:40] for cell in row.cells[:3]]
            print(f"Row {row_idx}: {' | '.join(cells)}")

    print("\n--- SPRAWDZENIE CZY PLACEHOLDERY ZOSTAŁY USUNIĘTE ---")
    placeholder_count = 0
    for para in doc.paragraphs:
        if '{' in para.text and '}' in para.text:
            # Pomijamy komentarz {Nie zmieniać}
            if 'Nie zmieniać' not in para.text:
                placeholder_count += 1
                print(f"  Pozostały placeholder: {para.text[:60]}...")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '{' in cell.text and '}' in cell.text:
                    placeholder_count += 1
                    print(f"  W tabeli: {cell.text[:60]}...")

    if placeholder_count == 0:
        print("  OK - Wszystkie placeholdery zostały wypełnione!")
    else:
        print(f"  UWAGA: Pozostało {placeholder_count} placeholderów")

    print("\n--- TABELA ZAŁĄCZNIKÓW ---")
    if len(doc.tables) > 1:
        table = doc.tables[1]
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip()[:30] for cell in row.cells]
            print(f"Row {row_idx}: {' | '.join(cells)}")


if __name__ == "__main__":
    main()
